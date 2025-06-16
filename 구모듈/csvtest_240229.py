import csv
from docx import Document


# CSV 파일 경로
file_path = "dam_data.csv"

damlist = {'감포댐':	1, '강정고령보':	2, '강천보':	3, '공주보':	4, '광동댐':	5, '구미보':	6, '구천댐':	7, '군남댐':	8, '군위댐':	9, '김천부항댐': 10,
           '낙단보':	11, '낙동강하굿둑': 12, '남강댐': 13, '달방댐':	14, '달성보': 15, '대곡댐':	16, '대암댐': 17, '대청댐': 18, '대청조정지': 19, '밀양댐': 20,
           '백제보댐':	21, '보령댐':	22, '보현산댐': 23, '부안댐':	24, '사연댐':	25, '상주보': 26, '선암댐':	27, '섬진강댐': 28, '성덕댐':	29, '세종보': 30,
           '소양강댐':	31, '수어댐':	32, '승촌보': 33, '안계댐':	34, '안동댐':	35, '안동조정지댐':	36, '여주보': 37, '연초댐':	38, '영주댐':	39, '영천댐':	40,
           '용담댐':	41, '운문댐':	42, '이포보': 43, '임하댐':	44, '임하조정지댐':	45, '장흥댐':	46, '주암댐': 47, '주암(조)': 48, '주암역조정지댐': 49, '죽산보': 50,
           '창녕함안보':	51, '충주댐':	52, '충주조정지댐':	53, '칠곡보': 54, '평림댐':	55, '평화의댐': 56, '한탄강댐': 57, '합천댐': 58, '합천조정지댐': 59, '합천창녕보': 60,
           '횡성댐': 61}

def namelist(name):
    try:
        row_index = damlist[name]
        print(row_index)
    except KeyError:
        print(f"The key '{name}' does not exist in the dictionary.")
        row_index = 0

    return row_index


def csvwrite(row_index):
    # 특정 행의 인덱스 (0부터 시작)
    target_row_index = row_index

    # 특정 행의 값을 담을 리스트
    row_values = None

    # CSV 파일을 열고 읽기 모드로 연다
    with open(file_path, newline='') as csvfile:
        # CSV 파일을 읽기 위한 reader 객체 생성
        reader = csv.reader(csvfile)

        # 각 행의 데이터를 읽어오기
        for index, row in enumerate(reader):
            # 특정 행을 찾았을 때
            if index == target_row_index:
                # 해당 행의 값을 리스트에 저장
                row_values = row
                break

    dam_name = row_values[1]    #댐이름
    dam_river = row_values[2]   #하천
    dam_type = row_values[3]    #형식
    dam_height = row_values[4]  #높이
    dam_length = row_values[5]  #길이
    dam_volume = row_values[6]  #체적
    dam_peak = row_values[7]    #정상표고
    dam_reservoirarea = row_values[8]    #유역면적
    dam_yearsupply = row_values[9]  #연간용수공급량
    dam_supplyarea = row_values[10] #저수면적
    dam_designfloodleve = row_values[11]    #계획홍수위
    dam_normalhighwaterlever = row_values[12]   #상시만수위
    dam_floodlimitlevel = row_values[13]    #홍수기제한수위
    dam_WRJpeak = row_values[14]    #월류정표고
    dam_lowwaterlevel = row_values[15]  #저수위
    dam_reservoirvolume= row_values[16] #총저수용량
    dam_realReservoirvolume = row_values[17]    #유효저수용량
    dam_floodcontrollevel = row_values[18]  #홍수조절용량
    dam_startday = row_values[19]   # 사업시작일
    dam_endday = row_values[20]    #사업종료일


    # 값이 담긴 리스트 출력
    for i in range(0,20):
        print(i, row_values[i])

    print(row_values)

    # print('0', row_values[0])
    # print('1', row_values[1])
    # print('2', row_values[2])
    # print('3', row_values[3])

    return row_values

if __name__ == '__main__':
    row_index = namelist('ㅁㅁㅁ')
    csvdata = csvwrite(row_index)

    doc = Document()
    doc.add_heading('CSV-Python 연동 시험 문서', level=1)
    doc.add_paragraph('이 문서는 csv 데이터를 읽어와 docx 문서로 작성하기 위하여 시험적으로 작성한 내용입니다.')
    doc.add_paragraph('0. 번호: ' + csvdata[0])
    doc.add_paragraph('1. 댐 이름: ' + csvdata[1])
    doc.add_paragraph('2. 하천: ' + csvdata[2])
    doc.add_paragraph('3. 형식: ' + csvdata[3])
    doc.add_paragraph('4. 높이: ' + csvdata[4])
    doc.add_paragraph('5. 길이: ' + csvdata[5])
    doc.add_paragraph('6. 길이: ' + csvdata[6])
    doc.add_paragraph('7. 정상표고: ' + csvdata[7])
    doc.add_paragraph('8. 정상표고: ' + csvdata[8])

    doc.save('CSV작성 시험_240305.docx')


