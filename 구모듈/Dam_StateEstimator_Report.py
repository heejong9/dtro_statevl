# Leakage = 누수
# Fail = 손상
# Spalling = 박리
# Desquamation = 층분리 \\ 박락
# Efflorescence = 백태
# Segregation = 재료 분리
# RebarExposure = 철근 노출
# Crack Repair = 균열(보수 후)

import configparser
# -*- coding: utf-8 -*-
import os
from datetime import datetime

import csv
import pandas as pd
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx2pdf import convert



# import docxtable



# import os

# parser = argparse.ArgumentParser(description="Read text from an image and filter by specific conditions")
# parser.add_argument('--measure-csv', required=True, help="검출 결과 measure 파일")
# parser.add_argument('--config', default="", help="시설물 Conf 파일")
# parser.add_argument('--stage4result', default="", help="상태평가4단계 결과")
# parser.add_argument('--output-path', default="", help="보고서 저장 경로")

# args = parser.parse_args()


# data = args.measure_csv
# data = pd.read_csv(data)

# data = pd.read_csv('.\\totalSheet.csv')
# data = pd.read_csv('.\\CJD_merge_test.csv')
# crackData = data[data['Type'] == 'Crack']
# defectData = data[data['Type'] != 'Crack']
#
# crackLen = len(crackData)
# defectLen = len(defectData)
# print(str(crackLen))
# print(str(defectLen))




# config = configparser.ConfigParser()
# config.read('.\\StateEstimateSetup.conf')
# config = config['Setup']
# TunnelType = config['TunnelType']
# TunnelName = config['TunnelName']
# Writer = config['Writer']
# SAFETY_DIAGNOSIS_DETAILED_GUIDELINE = config['안전진단세부지침년도']

###################

def gradeScore(grade) :

    if grade < 1.5 :
        return 'e'
    elif grade < 2.5 :
        return 'd'
    elif grade < 3.5:
        return 'c'
    elif grade < 4.5:
        return 'b'
    elif grade <= 5 :
        return 'a'

def adjustmentCoefficientScore(grade):

    if grade < 1.5 :
        return 6
    elif grade < 2.5 :
        return 6
    elif grade < 3.5:
        return 3
    elif grade < 4.5:
        return 2
    elif grade <= 5 :
        return 1

def createDirectory(directory):
    try:
        if not os.path.exists(directory):
            os.makedirs(directory)
    except OSError:
        print("Error: Failed to create the directory.")

    print(directory)

def Dam_StateEstimator_Report(config, file, outputpath, stateestimatorfiles, saveFile2, resultname):
    print(saveFile2)
    print('상태평가보고서 작성')
    print('config', os.path.dirname(config))
    print('config', config)

    print(resultname)

    data = pd.read_csv(file)

    # reportroot = 'StateEstimator'

    reportform = os.path.join(stateestimatorfiles, "form.docx")
    reportjson = config



    reportpicturepath = os.path.join(stateestimatorfiles, '그림')

    print('picturepath: ' + str(reportpicturepath))

    print('reportjson', reportjson)

    crackData = data[data['Type'] == 'Crack']
    defectData = data[data['Type'] != 'Crack']

    crackLen = len(crackData)
    defectLen = len(defectData)
    print(str(crackLen))
    print(str(defectLen))

    # config1 = config
    # config1 = config1

    # config1 = args.config

    # configs = os.join.path(reportroot, config)

    config1 = configparser.ConfigParser()
    config1.read(config, encoding='utf-8')

    config1 = config1['Cover']
    DamType = config1['DamType']
    DamName = config1['DamName']
    Writers = config1['Writer']
    now = datetime.now()

    #############################
    # document = Document('form.docx')
    document = Document(os.path.join(reportform))
    TitleName = DamName
    Today = str(now.year) + '-' + str(now.month) + '-' + str(now.day) #dt.second
    write = Writers
    document.add_heading(TitleName + ' 상태평가 보고서', 0)

    document.add_paragraph('\n\n\n\n\n\n\n\n\n\n\n\n\n\n\n\n\n')
    document.add_paragraph('작성일 : ' + Today, style='Today')
    document.add_paragraph('\n\n\n\n\n\n\n\n\n\n\n\n\n\n\n\n\n\n\n')
    document.add_paragraph('작성자 : ' + write, style='Write')

    document.add_paragraph('목   차', style='TitleName')
    document.add_paragraph('1. 시설물 개요', style='SubTitle')
    document.add_paragraph('1.1 시설물 현황 및 도면', style='TitleName2')
    document.add_paragraph('1.1.1 시설물 현황', style='subTitle2')
    document.add_paragraph('1.1.2 관련 도면', style='subTitle2')

    document.add_paragraph('2. 상태평가 개요\n', style='SubTitle')


    document.add_paragraph('3. 상태평가 기준 및 방법', style='SubTitle')
    document.add_paragraph('3.1 상태평가 항목 및 기준', style='TitleName2')
    document.add_paragraph('3.1.1 평가유형·영향계수 및 기준산정 방법', style='subTitle2')
    document.add_paragraph('3.1.2 상태평가 항목 및 기준', style='subTitle2')

    document.add_paragraph('3.2 상태평가 결과 산정 방법', style='TitleName2')
    document.add_paragraph('3.2.1 댐 시설물 평가 단계별 절차', style='subTitle2')
    document.add_paragraph('3.2.2 상태평가 단계별 구분', style='subTitle2')
    document.add_paragraph('3.2.3 기계 및 전기설비\n', style='subTitle2')

    document.add_paragraph('4. 상태평가 결과', style='SubTitle')


    document.add_page_break()
    document.add_paragraph(TitleName + ' 정밀안전진단 결과표', style='TitleName')
    document.add_paragraph('1. 기본현황', style='SubTitle')

    table = document.add_table(rows=11, cols=8)
    table.style = 'DefaultStyle'

    # config1 = configparser.ConfigParser()
    # config1.read('.\\DCDStateEstimateSetup_Table1.conf', encoding='utf-8')
    # config1 = config1['Table1']

    # config1.read(config, encoding='utf-8')
    # config1 = config1['Table1']

    Project_Name = config1['Project_Name']
    Project_Period = config1['Project_Period']
    Project_Manager = config1['Project_Manager']
    Project_Head = config1['Project_Head']
    Project_Contract = config1['Project_Contract']
    Project_Facility = config1['Project_Facility']
    Project_FacilityType = config1['Project_FacilityType']
    Project_FacilityClass = config1['Project_FacilityClass']
    Project_Completiondate = config1['Project_Completiondate']
    Project_Location = config1['Project_Location']
    Project_Scale = config1['Project_Scale']

    table.cell(0, 0).merge(table.cell(0, 7))
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = '가. 일반현황'

    hdr_cells = table.rows[1].cells
    table.cell(1, 1).merge(table.cell(1, 3))
    table.cell(1, 5).merge(table.cell(1, 7))
    hdr_cells[0].text = '용역명'
    hdr_cells[1].text = Project_Name
    hdr_cells[4].text = '진단기간'
    hdr_cells[5].text = Project_Period


    hdr_cells = table.rows[2].cells
    table.cell(2, 1).merge(table.cell(2, 3))
    table.cell(2, 5).merge(table.cell(2, 7))
    hdr_cells[0].text = '관리주체명'
    hdr_cells[1].text = Project_Manager
    hdr_cells[4].text = '대표자'
    hdr_cells[5].text = Project_Head


    hdr_cells = table.rows[3].cells
    table.cell(3, 1).merge(table.cell(3, 3))
    table.cell(3, 5).merge(table.cell(3, 7))
    hdr_cells[0].text = '공동수급'

    hdr_cells[4].text = '계약방법'
    hdr_cells[5].text = Project_Contract


    hdr_cells = table.rows[4].cells
    table.cell(4, 1).merge(table.cell(4, 3))
    hdr_cells[0].text = '시설물 구분'
    hdr_cells[1].text = Project_Facility

    hdr_cells[3].text = '종류'
    hdr_cells[4].text = Project_FacilityType

    hdr_cells[6].text = '종별'
    hdr_cells[7].text = Project_FacilityClass


    hdr_cells = table.rows[5].cells
    table.cell(5, 1).merge(table.cell(5, 3))
    hdr_cells[0].text = '준공일'
    hdr_cells[1].text = Project_Completiondate
    hdr_cells[4].text = '진단금액(천원)'
    hdr_cells[6].text = '안전등급'

    hdr_cells = table.rows[6].cells
    table.cell(6, 1).merge(table.cell(6, 3))
    table.cell(6, 5).merge(table.cell(6, 7))
    hdr_cells[0].text = '시설물 위치'
    hdr_cells[1].text = Project_Location
    hdr_cells[4].text = '시설물 규모'
    hdr_cells[1].text = Project_Scale


    hdr_cells = table.rows[7].cells
    table.cell(7, 0).merge(table.cell(7, 7))
    hdr_cells[0].text = '나. 진단 실시결과 현황'

    hdr_cells = table.rows[8].cells
    table.cell(8, 1).merge(table.cell(8, 7))
    hdr_cells[0].text = '중대결함'

    hdr_cells = table.rows[9].cells
    table.cell(9, 1).merge(table.cell(9, 7))
    hdr_cells[0].text = '진단 주요 결과'

    hdr_cells = table.rows[10].cells
    table.cell(10, 1).merge(table.cell(10, 7))
    hdr_cells[0].text = '주요 보수 보강'

    document.add_page_break()
    document.add_paragraph('2. 결과 요약', style='SubTitle')

    table = document.add_table(rows=2, cols=1)
    table.style = 'DefaultStyle'

    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = '책임 기술자 종합 의견'

    document.add_paragraph('\n\n\n\n 가. 정밀안전진단 외관조사 결과 기본사항', style='SubTitle')

    table = document.add_table(rows=9, cols=5)
    table.style = 'DefaultStyle'

    table.cell(0, 0).merge(table.cell(0, 4))
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = '상태평가 결과 및 보수 보강'
    hdr_cells[4].text = '종합평가 결과: '

    hdr_cells = table.rows[1].cells
    table.cell(1, 0).merge(table.cell(1, 1))
    hdr_cells[0].text = '결함발생 부재'
    hdr_cells[2].text = '상태평가 결과'
    hdr_cells[3].text = '결함 종류'
    hdr_cells[4].text = '보수 보강'

    hdr_cells = table.rows[2].cells
    hdr_cells[0].text = '필댐체'

    hdr_cells = table.rows[3].cells
    hdr_cells[0].text = '좌안비여수로'

    hdr_cells = table.rows[4].cells
    hdr_cells[0].text = '우안비여수로'

    hdr_cells = table.rows[5].cells
    table.cell(5, 0).merge(table.cell(8, 0))
    hdr_cells[0].text = '여수로'

    document.add_page_break()
    document.add_paragraph(TitleName + ' 현황표', style='TitleName')

    table = document.add_table(rows=13, cols=6)
    table.style = 'DefaultStyle'

    Facility_name = config1['Facility_name']
    Facility_managenumber = config1['Facility_managenumber']
    Facility_number = config1['Facility_number']
    Facility_Tel = config1['Facility_Tel']
    River_name = config1['River_name']
    DamSpecificType = config1['DamSpecificType']
    DamPeak = config1['DamPeak']
    DamHeight = config1['DamHeight']
    DamLength = config1['DamLength']
    DamVolume = config1['DamVolume']

    Reservoirvolume = config1['총저수량']
    realReservoirvolume = config1['유효저수량']
    Reservoirarea = config1['유역면적']
    DesignFloodLever = config1['계획홍수위']
    NormalHighWaterLever = config1['상시만수위']
    LowWaterLever = config1['저수위']

    SpillwayDesignFloodLever = config1['여수로계획홍수량/계획방류량']
    Spillwaygate = config1['여수로게이트']
    Spillwaytype = config1['여수로형식']

    SubSpillwayDesignFloodLever = config1['보조여수로계획홍수량/계획방류량']
    SubSpillwaygate = config1['보조여수로게이트']
    SubSpillwaytype = config1['보조여수로형식']



    hdr_cells = table.rows[0].cells
    table.cell(0, 1).merge(table.cell(0, 2))
    table.cell(0, 4).merge(table.cell(0, 5))
    hdr_cells[0].text = '시설물명'
    hdr_cells[1].text = Facility_name
    hdr_cells[3].text = '시설물번호'
    hdr_cells[4].text = Facility_number

    hdr_cells = table.rows[1].cells
    table.cell(1, 1).merge(table.cell(1, 2))
    table.cell(1, 4).merge(table.cell(1, 5))
    hdr_cells[0].text = '준공년월일'
    hdr_cells[1].text = Project_Completiondate

    hdr_cells[3].text = '관리번호'
    hdr_cells[4].text = Facility_managenumber

    hdr_cells = table.rows[2].cells
    table.cell(2, 1).merge(table.cell(2, 5))
    hdr_cells[0].text = '위치'
    hdr_cells[1].text = Project_Location


    hdr_cells = table.rows[3].cells
    table.cell(3, 1).merge(table.cell(3, 2))
    table.cell(3, 4).merge(table.cell(3, 5))
    hdr_cells[0].text = '관리주체'
    hdr_cells[1].text = Project_Manager
    hdr_cells[3].text = 'Tel.'
    hdr_cells[4].text = Facility_Tel


    hdr_cells = table.rows[4].cells
    table.cell(4, 0).merge(table.cell(9, 0))
    table.cell(4, 3).merge(table.cell(9, 3))
    hdr_cells[0].text = '댐제원'
    hdr_cells[1].text = '하천명'
    hdr_cells[2].text = River_name
    hdr_cells[3].text = '저수지제원'
    hdr_cells[4].text = '총저수량'
    hdr_cells[5].text = Reservoirvolume

    hdr_cells = table.rows[5].cells
    hdr_cells[1].text = '댐형식'
    hdr_cells[2].text = DamSpecificType
    hdr_cells[4].text = '유효저수량'
    hdr_cells[5].text = realReservoirvolume


    hdr_cells = table.rows[6].cells
    hdr_cells[1].text = '댐정상 표고'
    hdr_cells[2].text = DamPeak
    hdr_cells[4].text = '유역면적'
    hdr_cells[5].text = Reservoirarea


    hdr_cells = table.rows[7].cells
    hdr_cells[1].text = '댐 높이'
    hdr_cells[2].text = DamHeight
    hdr_cells[4].text = '계획 홍수위'
    hdr_cells[5].text = DesignFloodLever

    hdr_cells = table.rows[8].cells
    hdr_cells[1].text = '댐 길이'
    hdr_cells[2].text = DamLength
    hdr_cells[4].text = '상시 홍수위'
    hdr_cells[5].text = NormalHighWaterLever


    hdr_cells = table.rows[9].cells
    hdr_cells[1].text = '댐 체적'
    hdr_cells[2].text = DamVolume
    hdr_cells[4].text = '저수위'
    hdr_cells[5].text = LowWaterLever



    hdr_cells = table.rows[10].cells
    table.cell(10, 0).merge(table.cell(12, 0))
    table.cell(10, 3).merge(table.cell(12, 3))
    hdr_cells[0].text = '여수로'
    hdr_cells[1].text = '여수로계획홍수량/\n계획방류량'
    hdr_cells[2].text = SpillwayDesignFloodLever
    hdr_cells[3].text = '보조여수로'
    hdr_cells[4].text = '보조여수로계획홍수량/\n계획방류량'
    hdr_cells[5].text = SubSpillwayDesignFloodLever

    hdr_cells = table.rows[11].cells
    hdr_cells[1].text = '여수로게이트'
    hdr_cells[2].text = Spillwaygate
    hdr_cells[4].text = '보조여수로게이트'
    hdr_cells[5].text = SubSpillwaygate

    hdr_cells = table.rows[12].cells
    hdr_cells[1].text = '여수로형식'
    hdr_cells[2].text = Spillwaytype
    hdr_cells[4].text = '보조여수로형식'
    hdr_cells[5].text = SubSpillwaytype


    document.add_page_break()

    document.add_paragraph('1. 시설물 개요', style='TitleName')
    document.add_paragraph(Facility_name +'은 ' + Project_Location + '에 위치하고, '
                           '높이 ' + DamHeight + 'm, 길이 ' +DamLength + 'm의 ' + DamSpecificType + '형식으로 '+
                           Project_Completiondate + '년에 준공되었다. 댐의 제원 및 시설물의 현황은 다음과 같다.')

    document.add_paragraph('1.1 시설물 현황\n', style='SubTitle')

    document.add_paragraph('1.1.1. 시설물 현황', style='SubTitle')
    document.add_paragraph('   ' + TitleName + '의 기본 현황은 [표 2.1]과 같다.\n')

    document.add_paragraph('\n[표 2.1] 시설물 현황', style='tableName')

    document.add_paragraph(TitleName+' 현황표', style='TitleName')

    table = document.add_table(rows=13, cols=6)
    table.style = 'DefaultStyle'

    Facility_name = config1['Facility_name']
    Facility_managenumber = config1['Facility_managenumber']
    Facility_number = config1['Facility_number']
    Facility_Tel = config1['Facility_Tel']
    River_name = config1['River_name']
    DamSpecificType = config1['DamSpecificType']
    DamPeak = config1['DamPeak']
    DamHeight = config1['DamHeight']
    DamLength = config1['DamLength']
    DamVolume = config1['DamVolume']

    hdr_cells = table.rows[0].cells
    table.cell(0, 1).merge(table.cell(0, 2))
    table.cell(0, 4).merge(table.cell(0, 5))
    hdr_cells[0].text = '시설물명'
    hdr_cells[1].text = Facility_name
    hdr_cells[3].text = '시설물번호'
    hdr_cells[4].text = Facility_number

    hdr_cells = table.rows[1].cells
    table.cell(1, 1).merge(table.cell(1, 2))
    table.cell(1, 4).merge(table.cell(1, 5))
    hdr_cells[0].text = '준공년월일'
    hdr_cells[1].text = Project_Completiondate

    hdr_cells[3].text = '관리번호'
    hdr_cells[4].text = Facility_managenumber

    hdr_cells = table.rows[2].cells
    table.cell(2, 1).merge(table.cell(2, 5))
    hdr_cells[0].text = '위치'
    hdr_cells[1].text = Project_Location

    hdr_cells = table.rows[3].cells
    table.cell(3, 1).merge(table.cell(3, 2))
    table.cell(3, 4).merge(table.cell(3, 5))
    hdr_cells[0].text = '관리주체'
    hdr_cells[1].text = Project_Manager
    hdr_cells[3].text = 'Tel.'
    hdr_cells[4].text = Facility_Tel

    hdr_cells = table.rows[4].cells
    table.cell(4, 0).merge(table.cell(9, 0))
    table.cell(4, 3).merge(table.cell(9, 3))
    hdr_cells[0].text = '댐제원'
    hdr_cells[1].text = '하천명'
    hdr_cells[2].text = River_name
    hdr_cells[3].text = '저수지제원'
    hdr_cells[4].text = '총저수량'
    hdr_cells[5].text = Reservoirvolume


    hdr_cells = table.rows[5].cells
    hdr_cells[1].text = '댐형식'
    hdr_cells[2].text = DamSpecificType
    hdr_cells[4].text = '유효저수량'
    hdr_cells[5].text = realReservoirvolume


    hdr_cells = table.rows[6].cells
    hdr_cells[1].text = '댐정상 표고'
    hdr_cells[2].text = DamPeak
    hdr_cells[4].text = '유역면적'
    hdr_cells[5].text = Reservoirarea


    hdr_cells = table.rows[7].cells
    hdr_cells[1].text = '댐 높이'
    hdr_cells[2].text = DamHeight
    hdr_cells[4].text = '계획 홍수위'
    hdr_cells[5].text = DesignFloodLever


    hdr_cells = table.rows[8].cells
    hdr_cells[1].text = '댐 길이'
    hdr_cells[2].text = DamLength
    hdr_cells[4].text = '상시 홍수위'
    hdr_cells[5].text = NormalHighWaterLever


    hdr_cells = table.rows[9].cells
    hdr_cells[1].text = '댐 체적'
    hdr_cells[2].text = DamVolume
    hdr_cells[4].text = '저수위'
    hdr_cells[5].text = LowWaterLever



    hdr_cells = table.rows[10].cells
    table.cell(10, 0).merge(table.cell(12, 0))
    table.cell(10, 3).merge(table.cell(12, 3))
    hdr_cells[0].text = '여수로'
    hdr_cells[1].text = '여수로계획홍수량/\n계획방류량'
    hdr_cells[3].text = '보조여수로'
    hdr_cells[4].text = '보조여수로계획홍수량/\n계획방류량'

    hdr_cells = table.rows[11].cells
    hdr_cells[1].text = '여수로게이트'
    hdr_cells[4].text = '보조여수로게이트'

    hdr_cells = table.rows[12].cells
    hdr_cells[1].text = '여수로형식'
    hdr_cells[4].text = '보조여수로형식'


    document.add_page_break()

    document.add_paragraph('1.1.2. 관련도면', style='SubTitle')
    table = document.add_table(rows=1, cols=1)
    # tablePicture3 = '그림\\그림2_2.png'
    tablePicture3 = os.path.join(reportpicturepath, '그림2_2.png')

    table.style = 'DefaultStyle'
    hdr_cells = table.rows[0].cells
    paragraph = hdr_cells[0].paragraphs[0]

    run = paragraph.add_run()
    run.add_picture(os.path.join(reportpicturepath, '그림2_2.png'))


    document.add_paragraph('\n[그림2.1]' + Facility_name + ' 종평면도', style='tableName')

    table = document.add_table(rows=1, cols=1)
    # tablePicture5 = '그림\\그림2_3.png'
    tablePicture5 = os.path.join(reportpicturepath, '그림2_3.png')

    table.style = 'DefaultStyle'
    hdr_cells = table.rows[0].cells
    paragraph = hdr_cells[0].paragraphs[0]

    run = paragraph.add_run()
    run.add_picture(os.path.join(reportpicturepath, '그림2_3.png'))

    document.add_paragraph('\n[그림2.3]' + Facility_name + ' 표준 단면도', style='tableName')

    document.add_paragraph('2. 상태평가 개요 \n', style='TitleName')

    document.add_paragraph("시설물의 상태평가는 「시설물의 안전 및 유지관리 실시 지침(국토교통부, 국토안전관리원)」에 따라 실시하며, 상태평가에 대한 세부적인 사항은"
                           " 「시설물의 안전 및 유지관리 실시 세부지침(안전점검·진단편의 댐편, 2021.12)」을 준용하였다. "
                           "따라서, 「시설물의 안전 및 유지관리 실시 세부지침(안전점검·진단편의 댐편)」에 의거하여 외관조사 및 내구성 조사의 항목 및 수량에 따라 과업을 실시한 후, "
                           "상태평가를 위하여 중요 손상 및 결함을 세부 기준에 의해 분류하고 평가기법 및 절차에 따라 각 개별시설물에 대한 "
                           "결함의 등급과 점수 및 지수를 산정하여 상태등급을 최종적으로 결정하였다.")

    document.add_page_break()

    document.add_paragraph('3. 상태평가 기준 및 방법 \n', style='TitleName')
    document.add_paragraph('3.1 상태평가 항목 및 기준 \n ', style='SubTitle')
    document.add_paragraph('3.1.1 평가유형·영향계수 및 기준산정 방법 \n ', style='SubTitle')
    document.add_paragraph("시설물의 상태평가는 결함 및 손상에 따른 각각의 상태평가 기준을 적용하며, 상태변화가 전체 구조물에 미치는 안전성의 영향정도, "
                           "구조적인 중요도가 적절히 고려되어 평가될 수 있도록 결함 및 손상을 평가유형(評價類型)별로 구분하여 영향계수를 적용한다. \n"
                           "1) 평가유형의 구분 \n"
                           "결함 및 손상에 대한 평가유형은 다음과 같이 구분한다."
                           "\n① 중요결함"
                           "침하, 경사\\전도 및 활동 등과 같이 전체 구조물의 구조적인 안전에 직접 영향을 미치는 결함\n"
                           "② 국부결함\n"
                           "수평이음부 불량 등과 같이 구조물의 안전성에 직접적인 영향을 미치지는 않지만 손상이 진전될 경우 전체 구조물의 안전에 상당한 영향을 끼칠 수 있는 결함\n"
                           "③ 일반손상\n"
                           "파손, 마모, 콘크리트 재료분리 등과 같이 구조물의 안전에 크게 영향을 주지 않는 일반적인 손상\n"
                           "2) 영향계수의 적용\n"
                           "각 부재에서 발생하는 각종 손상 및 결함에 대한 상태평가 시 손상이 전체 구조물에 미치는 안전성의 영향정도, 구조적인 중요도가 적절히 고려되어 평가될 수 있도록 영향계수를 적용한다."
                           "영향계수는 안전성에 직접적인 영향을 미치는 중요 결함의 상태등급을 기준으로 하여 국부적인 결함의 등급을 상향 조정함으로써"
                           " 이들이 전체 구조물에 미치는 영향을 평가 절하하는 계수이며, 영향계수는 상태평가를 위한 표준기준이며, 책임기술자의 판단으로 다소 조정할 수 있다.""")
    document.add_page_break()

    document.add_paragraph('3.1.2 상태평가 항목 및 기준 \n ', style='SubTitle')


    document.add_paragraph("본 과업대상 시설물인"+"["+DamName+"]은"+"["+DamType+"]형식으로"+"["+DamType+"]에 준하는 기준을 적용하였다.\n"
                           "「세부지침」에 준하여 정량적이고 객관적인 상태평가를 위하여 부재별, 개별부재별, 복합부재별, 개별시설별 각 부재별 상태평가 항목은 다음과 같다.")

    document.add_paragraph('1) 콘크리트댐 \n ', style='SubTitle')
    document.add_paragraph('(가) 댐체 \n ', style='SubTitle')
    document.add_paragraph('① 댐마루 \n ', style='SubTitle')
    document.add_paragraph('\n[표3.1] 댐마루 상태평가 항목 및 기준', style='tableName')

    picture1 = os.path.join(reportpicturepath, '표3_1.png')
    document.add_picture(os.path.join(reportpicturepath, '표3_1.png'))


    # picture1 = document.add_picture('그림\\표3_1.png')
    last_paragraph = document.paragraphs[-1]
    last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    document.add_page_break()

    document.add_paragraph('② 상류면 \n ', style='SubTitle')
    document.add_paragraph('\n[표3.2] 상류면 상태평가 항목 및 기준', style='tableName')
    picture1 = os.path.join(reportpicturepath, '표3_2.png')
    document.add_picture(os.path.join(reportpicturepath, '표3_2.png'))

    # picture1 = document.add_picture('그림\\표3_2.png')
    last_paragraph = document.paragraphs[-1]
    last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    document.add_page_break()

    document.add_paragraph('③ 하류면 \n ', style='SubTitle')
    document.add_paragraph('\n[표3.3] 하류면 상태평가 항목 및 기준', style='tableName')
    picture1 = os.path.join(reportpicturepath, '표3_3.png')
    document.add_picture(os.path.join(reportpicturepath, '표3_3.png'))

    # picture1 = document.add_picture('그림\\표3_3.png')
    last_paragraph = document.paragraphs[-1]
    last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    document.add_page_break()


    document.add_paragraph('(나) 검사랑(갤러리) \n ', style='SubTitle')
    document.add_paragraph('\n[표3.4] 검사랑(갤러리) 상태평가 항목 및 기준', style='tableName')
    picture1 = os.path.join(reportpicturepath, '표3_4.png')
    document.add_picture(picture1)

    # picture1 = document.add_picture('그림\\표3_4.png')
    last_paragraph = document.paragraphs[-1]
    last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    document.add_page_break()


    document.add_paragraph('(다) 배수구 및 그라우팅 터널 \n ', style='SubTitle')
    document.add_paragraph('\n[표3.5] 배수구 및 그라우팅 터널 상태평가 항목 및 기준', style='tableName')
    picture1 = os.path.join(reportpicturepath, '표3_5.png')
    document.add_picture(os.path.join(reportpicturepath, '표3_5.png'))

    # picture1 = os.path.join(reportpicturepath, '표3_5.png')
    # picture1 = document.add_picture(picture1)


    last_paragraph = document.paragraphs[-1]
    last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    document.add_page_break()


    document.add_paragraph('(라) 기초 및 양안부 \n ', style='SubTitle')
    document.add_paragraph('\n[표3.6] 기초 및 양안부 상태평가 항목 및 기준', style='tableName')
    # picture1 = os.path.join(reportpicturepath, '표3_6.png')
    document.add_picture(os.path.join(reportpicturepath, '표3_6.png'))

    # picture1 = document.add_picture('그림\\표3_6.png')
    last_paragraph = document.paragraphs[-1]
    last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    document.add_page_break()


    document.add_paragraph('2) 일반적인 콘크리트 구조물 \n ', style='SubTitle')
    document.add_paragraph('\n[표3.7] 일반적인 콘크리트 구조물 상태평가 항목 및 기준', style='tableName')
    picture1 = os.path.join(reportpicturepath, '표3_7.png')
    document.add_picture(picture1)
    # picture1 = document.add_picture('그림\\표3_7.png')


    last_paragraph = document.paragraphs[-1]
    last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    document.add_page_break()


    document.add_paragraph('\n[표3.7] 일반적인 콘크리트 구조물 상태평가 항목 및 기준(계속)', style='tableName')
    picture1 = os.path.join(reportpicturepath, '표3_7_2.png')
    document.add_picture(os.path.join(reportpicturepath, '표3_7_2.png'))


    # picture1 = document.add_picture('그림\\표3_7_2.png')
    last_paragraph = document.paragraphs[-1]
    last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    document.add_page_break()

    document.add_paragraph('\n[표3.7] 일반적인 콘크리트 구조물 상태평가 항목 및 기준(계속)', style='tableName')
    # picture1 = os.path.join(reportpicturepath, '표3_7_3.png')
    document.add_picture(os.path.join(reportpicturepath, '표3_7_3.png'))

    # picture1 = document.add_picture('그림\\표3_7_3.png')
    last_paragraph = document.paragraphs[-1]
    last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    document.add_page_break()


    document.add_paragraph('3)기계설비 \n ', style='SubTitle')
    document.add_paragraph('(가) 권양기 \n ', style='SubTitle')

    document.add_paragraph('\n[표3.8] 기계설비-권양기 상태평가 항목 및 기준', style='tableName')
    # picture1 = os.path.join(reportpicturepath, '표3_8.png')
    document.add_picture(os.path.join(reportpicturepath, '표3_8.png'))

    # picture1 = document.add_picture('그림\\표3_8.png')
    last_paragraph = document.paragraphs[-1]
    last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER


    document.add_page_break()

    document.add_paragraph('(나) 수문 및 문틀 \n ', style='SubTitle')

    document.add_paragraph('\n[표3.9] 기계설비-수문 및 문틀 상태평가 항목 및 기준', style='tableName')
    # picture1 = os.path.join(reportpicturepath, '표3_9.png')
    document.add_picture(os.path.join(reportpicturepath, '표3_9.png'))
    # picture1 = document.add_picture('그림\\표3_9.png')
    last_paragraph = document.paragraphs[-1]
    last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    document.add_page_break()

    document.add_paragraph('4) 전기설비 \n ', style='SubTitle')

    document.add_paragraph('\n[표3.10] 전기설비 상태평가 항목 및 기준', style='tableName')
    # picture1 = os.path.join(reportpicturepath, '표3_10.png')
    document.add_picture(os.path.join(reportpicturepath, '표3_10.png'))


    # picture1 = document.add_picture('그림\\표3_10.png')
    last_paragraph = document.paragraphs[-1]
    last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    document.add_page_break()


    document.add_paragraph('3.2상태평가 결과 산정 방법 \n ', style='SubTitle')

    document.add_paragraph('3.2.1 댐 시설물 평가 단계별 절차 \n ', style='SubTitle')

    document.add_paragraph("댐 시설물에 대한 상태평가는 [그림 3.1]과 같이 단계별로 구분할 때 댐 시설물은 통합시설물 (6단계) 에 해당하는 시설물로서 간주하고, 하위단계인 복합시설, 개별시설, 복합부재, 개별부재로 구분한다."
                           "외관조사망도는 개별부재에 대하여 작성하는 것을 원칙으로 하고 필요시 개별부재의 크기, 면적에 따라 부위별로 분할하여 작성한다.")


    # picture1 = document.add_picture('그림\\그림3_1.png')
    document.add_picture(os.path.join(reportpicturepath, '그림3_1.png'))
    last_paragraph = document.paragraphs[-1]
    last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    document.add_paragraph('\n[그림3.1] 댐 시설물 평가 단계별 절차', style='tableName')

    document.add_page_break()

    document.add_paragraph('3.2.2 상태평가 단계별 구분 \n ', style='SubTitle')

    document.add_paragraph("시설물의 상태를 평가하기 위하여 시설물을 단계별로 구분하여 다음 표와 같이 평가단계별 구분표를 작성하고 본 보고서에 수록한다.")

    document.add_paragraph('\n[표3.11] 댐 시설물의 상태평가 단계별 구분표(예시)', style='tableName')
    # picture1 = document.add_picture('그림\\표3_11.png')
    document.add_picture(os.path.join(reportpicturepath, '표3_11.png'))

    last_paragraph = document.paragraphs[-1]
    last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    document.add_page_break()

    document.add_paragraph('1) 1단계 상태평가 : 부재(部材)별 손상상태 평가표 작성 \n ', style='SubTitle')
    document.add_paragraph("시설물의 상태평가 단계별 구분표에 따라 개별부재를 1개 외관조사망도 또는 필요에 따라 부위별로"
                           " 다수의 외관조사망도로 구분하여 개략도에 손상 및 결함상태를 도시하고,"
                           " 조사결과표에 개별부재에 대한 손상내용을 상세히 기록한 후, 그 손상 정도에 대하여 5단계(a～e) 상태평가 결과 및 평가점수를 부여한다.\n"
                           "○  손상상태 평가표에는 평가항목에 없는 상태변화라 할지라도 모두 기록하는 것을 원칙으로 한다.\n"
                           "○ 각 상태변화에 대한 상태평가 결과가 c, d, e 등급일 경우 보수ㆍ보강 우선순위에 따라 보수ㆍ보강을 한다.\n")

    document.add_paragraph('\n[표3.12] 부재(부위)별 손상상태 평가표(1단계) (예)', style='tableName')
    # picture1 = document.add_picture('그림\\표3_12.png')
    document.add_picture(os.path.join(reportpicturepath, '표3_12.png'))

    last_paragraph = document.paragraphs[-1]
    last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    document.add_page_break()

    document.add_paragraph('2) 2단계 상태평가 : 개별부재(個別部材) 평가표 작성  \n ', style='SubTitle')
    document.add_paragraph("댐체, 여수로 및 수로터널 등과 같이 길거나 또는 면적이 넓은 슬래브는 이를 1개의 개별부재로"
                           " 평가할 경우 일부에 발생한 손상이 평가결과에 미치는 영향이 크므로 콘크리트 구조물에서는"
                           " 그 손상이 부재에 영향을 미칠 수 있는 범위(길이 10～20m) 또는 수축이음부, 제체에서는 수십～수백m로 적절히 동일 규모가 되도록 분할하여 각각을 개별부재로서 평가한다."
                           "\n○ 개별 부재별로 작성된 외관조사망도에 나타난 손상 및 결함을 평가유형별로 중요결함, 국부결함, 일반손상으로 구분한다."
                           "\n○ 개별부재의 평가는 각각의 손상 및 결함에 대한 평가기준에 따른 평가점수(M)에 손상 및 결함이 부재의 안전에 미치는 영향을 반영한 평가유형별 영향계수(F)를 곱하여 산출한다."
                           "\n○ 산출된 결함 및 손상의 상태평가지수(E1) 중 최솟값을 개별부재의 상태평가지수(E2) 및 상태평가 결과를 결정한다")

    document.add_page_break()

    document.add_paragraph('\n[표3.13] 상태평가 결과별 평가지수 및 평가유형별 영향 계수', style='tableName')
    # picture1 = document.add_picture('그림\\표3_13.png')
    document.add_picture(os.path.join(reportpicturepath, '표3_13.png'))

    last_paragraph = document.paragraphs[-1]
    last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER


    document.add_paragraph('\n[표3.14] 개별부재 평가표(2단계) (예)', style='tableName')
    # picture1 = document.add_picture('그림\\표3_13.png')
    document.add_picture(os.path.join(reportpicturepath, '표3_14.png'))

    last_paragraph = document.paragraphs[-1]
    last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER


    document.add_page_break()

    document.add_paragraph('3) 3단계 상태평가 : 복합부재(複合部材) 평가표 작성  \n ', style='SubTitle')
    document.add_paragraph("○ 복합부재는 개별부재의 집합으로 주요부재와 보조부재로 구분할 수 있다."
                           "\n○ 복합부재의 평가는 개별부재가 복합부재의 안전에 미치는 영향을 판단하여 그 중요도를 반영한다. 콘크리트 부재는 조사망 면적 비율을 기준으로 중요도를 결정한다. 이때 개별부재의 중요도의 합이 100이 되도록 규정한다."
                           "\n○ 중요도가 규정되지 않은 추가적인 개별부재가 있는 경우에는 그 개별부재의 중요도를 판단하여 정하고, 기타의 부재들은 규정된 비율대로 배분한다."
                           "\n○ 책임기술자는 개별부재의 특성에 따라 중요도를 조정할 필요가 있다고 판단될 경우 규정된 값의 20%값 범위 내에서 조정할 수 있다."
                           "\n○ 또한, 복합부재의 안전은 상태가 나쁜 개별부재의 영향을 크게 받으므로 그에 상응한 보정을 하기 위하여 조정계수를 사용한다."
                           "\n○ 복합부재의 평가지수(E3) 산정 시 조정계수의 사용은 개별부재의 평가지수(E2)별로 위험성이 큰 값에 보다 큰 가중치를 적용하여 부재 전체의 안전성을 평가 절하한다."
                           "이는 단순 산술평균법의 적용보다 다소 낮은 평가지수의 평가결과를 도출한다."
                           "\n○ 복합부재의 평가는 개별부재의 평가지수(E2)에 중요도 및 조정계수를 반영하여 복합부재의 상태평가지수(E3)를 산출하고 상태평가 결과를 결정한다."
                           "\n▷ 복합부재의 상태평가지수(E3) = ∑(E2 × A × W) \\ ∑(A × W)"
                           "\n여기서,	E2	: 개별부재의 상태평가지수"
                           "\nA	: 조정계수"
                           "\nW	: 중요도")


    document.add_page_break()
    document.add_paragraph('\n[표3.15] 평가지수에 따른 조정계수', style='tableName')
    # picture1 = document.add_picture('그림\\표3_15.png')
    document.add_picture(os.path.join(reportpicturepath, '표3_15.png'))

    last_paragraph = document.paragraphs[-1]
    last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    document.add_paragraph('\n[표3.16] 중요도 조정방법 (예)', style='tableName')
    # picture1 = document.add_picture('그림\\표3_16.png')
    document.add_picture(os.path.join(reportpicturepath, '표3_16.png'))

    last_paragraph = document.paragraphs[-1]
    last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    document.add_paragraph("◦ 상기 예시는 시설물에서 어느 특정 부재가 추가되거나, 없는 경우에 중요도를 조정하여 중요도의 합이 100이 되도록 조정하기 위한 방법이다.")

    document.add_paragraph('\n[표3.17] 복합부재 평가표(3단계) (예)', style='tableName')
    # picture1 = document.add_picture('그림\\표3_17.png')
    document.add_picture(os.path.join(reportpicturepath, '표3_17.png'))

    last_paragraph = document.paragraphs[-1]
    last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    document.add_page_break()

    document.add_paragraph('4) 4단계 상태평가 : 개별시설(個別施設) 평가표 작성   \n ', style='SubTitle')
    document.add_paragraph("○ 댐의 제체는 개별시설로서 동일기능을 수행하는 복합부재(블록1, 블록2, …)의 집합으로 구성되어 있다."
                           "\n○ 개별시설의 평가는 복합부재의 중요도는 같다는 가정하에 복합부재의 상태평가지수(E3)에 규모"
                           "(길이, 면적, 부피, Capacity 등)를 반영하여 개별시설의 상태평가지수(Ec)를 산출하고 상태평가 결과를 결정한다."
                           " 댐 시설물에서 규모(S)값은 조사망 면적(㎡)을 사용하는 것을 원칙으로 하고, 책임기술자의 판단에 따라 길이, 부피 등을 사용할 수 있다."
                           "\n○ 또한 개별시설의 평가단계에서는 안전성평가를 수행하여 종합평가 결과를 결정한다."
                           "\n▷ 개별시설의 상태평가지수(Ec) = Min + V1 × V2"
                           "\n여기서, V1 = 0.3 × (Max - Min)"
                           "\nV2 = ∑(E3 × S) \\ (5 × ∑S)"
                           "\nS : 규모"
                           "\nMax : 복합부재의 상태평가지수(E3) 최댓값"
                           "\nMin : 복합부재의 상태평가지수(E3) 최솟값")

    document.add_paragraph('\n[표3.18] 개별시설 평가표(4단계) (예)', style='tableName')
    # picture1 = document.add_picture('그림\\표3_18.png')
    document.add_picture(os.path.join(reportpicturepath, '표3_18.png'))

    last_paragraph = document.paragraphs[-1]
    last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER


    document.add_page_break()

    document.add_paragraph('3.2.3 기계 및 전기설비 \n ', style='SubTitle')
    document.add_paragraph("○ 기전설비의 상태를 평가하기 위한 평가단계별 구분은 수문을 개별시설에 해당하는 것으로 하고, 수문 및 권양기, 전기설비로 구분하여 복합부재로 평가한다."
                           "\n○ 또한, 각각의 복합부재를 다음 표와 같이 개별부재로 분류하고, 설치되어 있는 개별부재의 중요도는 동일하게 적용한다."
                           "\n○ 4단계 평가시 규모는 복합부재의 중요도로써 정한다."
                           "\n복합부재의 중요도는 권양기 30%, 수문 40%, 전기설비 30%를 적용한다."
                           "\n○ 기전설비는 여수로, 취수시설 등 해당시설물의 개별시설로 평가하고, 토목시설과 함께 복합시설을 평가한다."
                           "\n○ 5단계 평가시 개별시설의 중요도는 토목시설 90%, 기전설비 10%를 적용한다."
                           "\n○ 책임기술자는 현장 여건에 따라 중요도를 20% 범위 내에서 조정할 수 있다."
                           "\n○ 기전설비의 손상상태평가표는 복합부재에 대하여 작성하며, 주로 손상상태를 기록하고 필요한 경우에만 개략도를 포함하여 작성한다."
                           "\n○ 기전설비의 상태평가 절차는 댐 시설물과 같은 방법 및 절차로 수행한다.")

    document.add_page_break()

    document.add_paragraph('\n[표3.19] 기계 및 전기설비의 상태평가 단계별 구분표', style='tableName')
    # picture1 = document.add_picture('그림\\표3_18.png')
    document.add_picture(os.path.join(reportpicturepath, '표3_19.png'))

    last_paragraph = document.paragraphs[-1]
    last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    document.add_paragraph("※) 개별부재(부위)에 대한 외관조사망도 작성")

    document.add_page_break()



    document.add_paragraph('4.상태평가 결과 \n ', style='SubTitle')

    document.add_paragraph('4.1 상태평가 4단계 결과 \n ', style='SubTitle')

    document.add_paragraph('\n[표 4.1] 4단계 결과', style='tableName')
    # document.add_picture(os.path.join(reportpicturepath, '표4_1.png'))

    # stage4data = args.stage4result
    stage4data = saveFile2
    stage4data = pd.read_csv(stage4data)
    # stage4data = pd.read_csv('4단계결과_231012_1.csv')

    FacilityList =[]
    EcList = []
    GradeList = []

    FacilityName = stage4data["StructureType"].copy()
    EcScore = stage4data["Ec"]
    GradeScore = stage4data["Grade"]

    table = document.add_table(rows=int(len(FacilityName)+1), cols=3)
    table.style = 'DefaultStyle'

    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = '개별시설'
    hdr_cells[1].text = '상태평가지수(Ec)'
    hdr_cells[2].text = '상태평가 결과'

    tableCount = 0

    print(len(FacilityName))
    print(FacilityName, EcScore, GradeScore)


    # for i in range(0, len(FacilityName)):
    #     FacilityList.append(str(FacilityName[i]))
    #     EcList.append(float(EcScore[i]))
    #     GradeList.append(str(GradeScore[i]))

    print(FacilityName[0])

    if FacilityName[0] == 'overflow':
        FacilityName[0] = '여수로'

    print('int(len(FacilityName))', int(len(FacilityName)))

    for i in range(0, int(len(FacilityName))):
        hdr_cells = table.rows[i+1].cells
        # hdr_cells = table.rows[i].cells
        hdr_cells[0].text = str(FacilityName[i])
        hdr_cells[1].text = str(EcScore[i])
        hdr_cells[2].text = str(GradeScore[i])
        tableCount += 1

    # with open(sys.argv[1],"r", encoding='cp949') as f :
    #     lines = f.readlines()
    #
    #     for i in lines :
    #         line = i.replace('\n', "")
    #
    #         document.add_paragraph(line)


    # document.add_paragraph('댐 시설물에 대한 전체 균열 수는 ['+ str(crackLen)  + ']개, 결함 수는 ['+  str(defectLen) + ']개로 확인되었다.')

    # document.add_heading(TitleName + ' 상태평가 보고서', 0)
    document.add_paragraph('4.2 손상물량표 \n ', style='SubTitle')
    document.add_paragraph(TitleName +' 결함의 대한 손상 물량은 아래와 같다.')
    with open(os.path.join(outputpath, 'merged_file.csv'), "r", encoding="cp949", errors="replace",
              newline="") as csv_file:
        csv_reader = csv.reader(csv_file)

        # 먼저 CSV 파일의 내용을 리스트로 변환
        data = list(csv_reader)

        if data:
            # 최대 열 수 계산
            max_columns = max(len(row) for row in data)
            #
            # # DOCX 표 생성
            table = document.add_table(rows=1, cols=max_columns)
            table.style = 'DefaultStyle'
            #
            # # 첫 번째 행에 표 제목 추가
            table.cell(0, 0).merge(table.cell(0, 4))
            table.rows[0].cells[0].text = "손상물량표"

            # CSV 파일 내용을 DOCX 표로 복사
            for row in data:
                cells = table.add_row().cells
                for i, cell_value in enumerate(row):
                    cells[i].text = cell_value

            table.cell(2, 0).merge(table.cell(14, 0))
            table.cell(15, 0).merge(table.cell(27, 0))
            table.cell(28, 0).merge(table.cell(40, 0))
            table.cell(41, 0).merge(table.cell(53, 0))


            table.cell(2, 1).merge(table.cell(7, 1))
            table.cell(15, 1).merge(table.cell(21, 1))
            table.cell(28, 1).merge(table.cell(34, 1))
            table.cell(41, 0).merge(table.cell(47, 0))


        else:
            # 빈 파일 처리 (선택 사항)
            document.add_paragraph("CSV 파일이 비어 있습니다.")
    # 저장
    # document.save(str(outputpath) + '\\demo_231017.docx')
    # file = open(str(outputpath) + '\\demo_231017.pdf', "w")
    # file.close()
    # convert(str(outputpath) + '\\demo_231017.docx', str(outputpath) + '\\demo_231017.pdf')

    document.save(str(outputpath) + '\\'+resultname+'.docx')
    # file = open(str(outputpath) + '\\'+reportname+'.pdf', "w")
    # file.close()
    convert(str(outputpath) + '\\'+resultname+'.docx', str(outputpath) + '\\'+resultname+'.pdf')
#
# if __name__ == '__main__':
#     # Dam_StateEstimator_Report('DCDStateEstimateSetup.conf', 'CJD_merge_test2.csv', 'output', 'output\\4단계결과_231016.csv')
#     # Dam_StateEstimator_Report.Dam_StateEstimator_Report(config, file, outputpath, saveFile2 , '상태평가보고서_Dam_'+Today)
#     conf = 'D:/우림/사내보고서/2023년 작업/SaaS/workspace/01. Dam/Dam_Fill_231123/StateEstimatorReport/SSYDStateEstimateSetup.conf'
#     Dam_StateEstimator_Report('SYDStateEstimateSetup.conf', 'SYD.csv', 'test1123', './output_231123/4단계결과_2023-11-23.csv', '상태평가보고서_Dam_231123')
