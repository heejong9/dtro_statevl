import csv
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches

# CSV 파일 경로와 DOCX 파일 이름 지정
csv_file_path = "merged_file.csv"
docx_file_name = "output_document.docx"

# DOCX 문서 생성
doc = Document()

document = Document('form.docx')


# CSV 파일을 열고 내용을 DOCX 표로 추가
with open(csv_file_path, "r", encoding="cp949", errors="replace", newline="") as csv_file:
    csv_reader = csv.reader(csv_file)

    # 먼저 CSV 파일의 내용을 리스트로 변환
    data = list(csv_reader)

    if data:
        # 최대 열 수 계산
        max_columns = max(len(row) for row in data)

        # DOCX 표 생성
        table = document.add_table(rows=1, cols=max_columns)
        table.style = 'DefaultStyle'

        # 첫 번째 행에 표 제목 추가
        table.cell(0,0).merge(table.cell(0, 4))
        table.rows[0].cells[0].text = "손상물량표"

        # CSV 파일 내용을 DOCX 표로 복사
        for row in data:
            cells = table.add_row().cells
            for i, cell_value in enumerate(row):
                cells[i].text = cell_value
    else:
        # 빈 파일 처리 (선택 사항)
        document.add_paragraph("CSV 파일이 비어 있습니다.")

# DOCX 파일 저장
document.save(docx_file_name)

print(f"{csv_file_path} 내용이 {docx_file_name}에 성공적으로 추가되었습니다.")

